from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Optional, Tuple

from docx import Document
from pypdf import PdfReader


@dataclass
class LoadedDocument:
    filename: str
    doc_type: str  # "pdf" | "docx" | "txt"
    text: str


def _load_pdf_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages_text = []
    for page in reader.pages:
        pages_text.append(page.extract_text() or "")
    return "\n".join(pages_text).strip()


def _load_docx_text(content: bytes) -> str:
    doc = Document(BytesIO(content))
    paras = [p.text for p in doc.paragraphs if p.text is not None]
    return "\n".join(paras).strip()


def _load_txt_text(content: bytes) -> str:
    return content.decode("utf-8", errors="replace").strip()


def load_text_from_upload(filename: str, content: bytes) -> LoadedDocument:
    """
    Extracts text from PDF (text-based), DOCX, or TXT.
    Raises ValueError for unsupported types or empty extraction.
    """
    lower = filename.lower().strip()

    if lower.endswith(".pdf"):
        doc_type = "pdf"
        text = _load_pdf_text(content)
    elif lower.endswith(".docx"):
        doc_type = "docx"
        text = _load_docx_text(content)
    elif lower.endswith(".txt"):
        doc_type = "txt"
        text = _load_txt_text(content)
    else:
        raise ValueError("Unsupported file type. Upload a PDF (text-based), DOCX, or TXT.")

    # Minimal validation
    if not text or len(text.strip()) < 20:
        raise ValueError(
            "Could not extract usable text. If PDF is scanned, convert to text-based PDF or use DOCX/TXT."
        )

    return LoadedDocument(filename=filename, doc_type=doc_type, text=text)
